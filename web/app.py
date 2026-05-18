"""
RAG Web Application - FastAPI Backend (v3 — Upgraded Models)
Pipeline: Upload PDF/Word → Read & Chunk → Embedding → ChromaDB → Query with LLM

Models:
  - Embedding: BAAI/bge-m3 (multilingual, 1024-dim, top-tier retrieval)
  - LLM: Qwen/Qwen2.5-7B-Instruct (7B params, excellent reasoning)

GPU L4 (24GB VRAM): bge-m3 ~2GB + Qwen2.5-7B fp16 ~14GB = ~16GB total ✓
"""

import os
import re
import uuid
import logging
from contextlib import asynccontextmanager
from pathlib import Path

import torch
import chromadb
from sentence_transformers import SentenceTransformer
from transformers import AutoTokenizer, AutoModelForCausalLM

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# ============================================================
# Config
# ============================================================
UPLOAD_FOLDER = "uploads"
CHROMA_DIR = "chroma_db"
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}

# --- Chunking ---
CHUNK_SIZE = 512          # tokens xấp xỉ, tăng lên vì bge-m3 hỗ trợ 8192 tokens
CHUNK_OVERLAP = 64
TOP_K = 5

# --- Models (UPGRADED) ---
EMBED_MODEL_NAME = "BAAI/bge-m3"          # 568M params, 1024-dim, multilingual
LLM_MODEL_NAME = "Qwen/Qwen2.5-7B-Instruct"  # 7B params, excellent quality

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CHROMA_DIR, exist_ok=True)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ============================================================
# Pydantic schemas
# ============================================================
class QueryRequest(BaseModel):
    session_id: str
    question: str


class QueryResponse(BaseModel):
    answer: str
    sources: list[str] = []


class UploadResponse(BaseModel):
    success: bool
    session_id: str
    filename: str
    chunk_count: int
    message: str


class SessionInfo(BaseModel):
    session_id: str
    filename: str
    chunk_count: int


class HealthResponse(BaseModel):
    status: str
    gpu: bool
    gpu_name: str
    vram_used_gb: float
    vram_total_gb: float
    embed_model: str
    llm_model: str


# ============================================================
# Model manager — load once, reuse everywhere
# ============================================================
class ModelManager:
    def __init__(self):
        self.embedding_model: SentenceTransformer | None = None
        self.tokenizer = None
        self.llm = None
        self.chroma_client: chromadb.PersistentClient | None = None

    def load_all(self):
        """Pre-load all models at startup."""
        # 1. Embedding model — bge-m3
        logger.info(f"[1/3] Loading embedding: {EMBED_MODEL_NAME} ...")
        self.embedding_model = SentenceTransformer(EMBED_MODEL_NAME)
        logger.info(f"  ✓ Embedding loaded (dim={self.embedding_model.get_sentence_embedding_dimension()})")

        # 2. LLM — Qwen2.5-7B-Instruct
        logger.info(f"[2/3] Loading LLM: {LLM_MODEL_NAME} ...")
        self.tokenizer = AutoTokenizer.from_pretrained(
            LLM_MODEL_NAME, trust_remote_code=True
        )
        self.llm = AutoModelForCausalLM.from_pretrained(
            LLM_MODEL_NAME,
            torch_dtype=torch.float16,
            device_map="auto",
            trust_remote_code=True,
        )
        logger.info("  ✓ LLM loaded")

        # Log VRAM usage
        if torch.cuda.is_available():
            used = torch.cuda.memory_allocated() / 1e9
            total = torch.cuda.get_device_properties(0).total_memory / 1e9
            logger.info(f"  VRAM: {used:.1f}GB / {total:.1f}GB")

        # 3. ChromaDB
        self.chroma_client = chromadb.PersistentClient(path=CHROMA_DIR)
        logger.info("[3/3] ChromaDB ready ✓")


models = ModelManager()

# Store sessions
sessions: dict[str, dict] = {}


# ============================================================
# Lifespan
# ============================================================
@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("🚀 Starting up — loading models...")
    models.load_all()
    logger.info("✅ All models loaded. Server ready!")
    yield
    logger.info("🛑 Shutting down...")


# ============================================================
# FastAPI app
# ============================================================
app = FastAPI(
    title="RAG Assistant API v3",
    description="Upload PDF/Word → Hỏi đáp với BGE-M3 + Qwen2.5-7B",
    version="3.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============================================================
# Document readers
# ============================================================
def read_pdf(filepath: str) -> str:
    """Đọc text từ PDF bằng PyMuPDF."""
    import fitz

    doc = fitz.open(filepath)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text


def read_docx(filepath: str) -> str:
    """Đọc text từ file Word (.docx) bằng python-docx."""
    from docx import Document

    doc = Document(filepath)
    parts = []

    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def read_document(filepath: str) -> str:
    """Đọc tài liệu — tự nhận dạng PDF hoặc Word."""
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return read_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return read_docx(filepath)
    else:
        raise ValueError(f"Không hỗ trợ định dạng: {ext}")


# ============================================================
# Chunking — sentence-aware splitting (tốt hơn cắt cứng)
# ============================================================
SENTENCE_SPLIT_RE = re.compile(r'(?<=[.!?。！？\n])\s+')


def chunk_text(
    text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP
) -> list[str]:
    """
    Chia văn bản thành các đoạn nhỏ, ưu tiên cắt ở ranh giới câu.
    Tránh cắt giữa câu → context mạch lạc hơn → LLM trả lời tốt hơn.
    """
    sentences = SENTENCE_SPLIT_RE.split(text)
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue

        if len(current_chunk) + len(sentence) + 1 <= chunk_size:
            current_chunk = f"{current_chunk} {sentence}".strip()
        else:
            if current_chunk:
                chunks.append(current_chunk)
            # Nếu 1 câu dài hơn chunk_size → cắt cứng
            if len(sentence) > chunk_size:
                start = 0
                while start < len(sentence):
                    chunks.append(sentence[start:start + chunk_size].strip())
                    start += chunk_size - overlap
            else:
                current_chunk = sentence

    if current_chunk:
        chunks.append(current_chunk)

    # Thêm overlap bằng cách gộp cuối chunk trước vào đầu chunk sau
    if overlap > 0 and len(chunks) > 1:
        overlapped = [chunks[0]]
        for i in range(1, len(chunks)):
            prev_tail = chunks[i - 1][-overlap:]
            overlapped.append(f"{prev_tail} {chunks[i]}".strip())
        chunks = overlapped

    return [c for c in chunks if c.strip()]


# ============================================================
# Embedding & Vector DB
# ============================================================
def embed_chunks(chunks: list[str]) -> list[list[float]]:
    """Tạo vector embeddings bằng BGE-M3."""
    embeddings = models.embedding_model.encode(
        chunks,
        show_progress_bar=False,
        normalize_embeddings=True,   # bge-m3 khuyến nghị normalize
    )
    return embeddings.tolist()


def store_to_chroma(
    collection_name: str, chunks: list[str], embeddings: list[list[float]]
) -> int:
    """Lưu chunks + embeddings vào ChromaDB."""
    try:
        models.chroma_client.delete_collection(collection_name)
    except Exception:
        pass

    collection = models.chroma_client.create_collection(
        name=collection_name,
        metadata={"hnsw:space": "cosine"},   # cosine similarity cho normalized vectors
    )
    ids = [f"chunk_{i}" for i in range(len(chunks))]
    collection.add(ids=ids, documents=chunks, embeddings=embeddings)
    return len(chunks)


def query_chroma(
    collection_name: str, query_text: str, top_k: int = TOP_K
) -> list[str]:
    """Tìm kiếm chunks liên quan trong ChromaDB."""
    collection = models.chroma_client.get_collection(name=collection_name)
    query_embedding = models.embedding_model.encode(
        [query_text],
        normalize_embeddings=True,
    ).tolist()
    results = collection.query(query_embeddings=query_embedding, n_results=top_k)
    return results["documents"][0] if results["documents"] else []


# ============================================================
# LLM Generation — improved prompt
# ============================================================
def generate_answer(query: str, context_chunks: list[str]) -> str:
    """Sinh câu trả lời bằng Qwen2.5-7B-Instruct dựa trên context."""
    context = "\n\n---\n\n".join(
        f"[Đoạn {i+1}] {chunk}" for i, chunk in enumerate(context_chunks)
    )

    messages = [
        {
            "role": "system",
            "content": (
                "Bạn là trợ lý AI chuyên phân tích tài liệu. "
                "Quy tắc:\n"
                "1. Chỉ trả lời dựa trên ngữ cảnh được cung cấp.\n"
                "2. Nếu không tìm thấy thông tin, nói rõ: 'Tài liệu không chứa thông tin này.'\n"
                "3. Trả lời bằng tiếng Việt, rõ ràng, có cấu trúc.\n"
                "4. Trích dẫn số đoạn [Đoạn X] khi tham chiếu thông tin."
            ),
        },
        {
            "role": "user",
            "content": f"Ngữ cảnh từ tài liệu:\n\n{context}\n\nCâu hỏi: {query}",
        },
    ]

    text_input = models.tokenizer.apply_chat_template(
        messages, tokenize=False, add_generation_prompt=True
    )
    inputs = models.tokenizer(text_input, return_tensors="pt").to(models.llm.device)

    with torch.no_grad():
        outputs = models.llm.generate(
            **inputs,
            max_new_tokens=2048,     # tăng từ 1024 → 2048, 7B sinh text dài tốt hơn
            temperature=0.7,
            top_p=0.9,
            do_sample=True,
            repetition_penalty=1.1,  # giảm lặp lại
        )

    response = models.tokenizer.decode(
        outputs[0][inputs["input_ids"].shape[1]:], skip_special_tokens=True
    )
    return response.strip()


# ============================================================
# API Routes
# ============================================================
@app.get("/", include_in_schema=False)
async def serve_index():
    return FileResponse("static/index.html")


@app.post("/api/upload", response_model=UploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """Upload PDF/Word → Đọc → Chunk → Embedding → Lưu ChromaDB."""
    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Chỉ chấp nhận file PDF hoặc Word (.docx). Bạn upload file {ext}",
        )

    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File quá lớn (tối đa 50MB).")

    try:
        session_id = str(uuid.uuid4())[:8]
        safe_name = file.filename.replace(" ", "_")
        filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}_{safe_name}")
        with open(filepath, "wb") as f:
            f.write(content)
        logger.info(f"Saved: {filepath}")

        # 1. Read
        logger.info(f"Step 1/4: Reading {ext} ...")
        text = read_document(filepath)
        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="Không đọc được nội dung. File có thể trống hoặc là ảnh scan.",
            )

        # 2. Chunk
        logger.info("Step 2/4: Chunking...")
        chunks = chunk_text(text)
        logger.info(f"  → {len(chunks)} chunks")

        # 3. Embed
        logger.info("Step 3/4: Embedding with BGE-M3...")
        embeddings = embed_chunks(chunks)

        # 4. Store
        logger.info("Step 4/4: Storing to ChromaDB...")
        collection_name = f"doc_{session_id}"
        chunk_count = store_to_chroma(collection_name, chunks, embeddings)

        sessions[session_id] = {
            "collection_name": collection_name,
            "filename": safe_name,
            "chunk_count": chunk_count,
            "filepath": filepath,
        }

        return UploadResponse(
            success=True,
            session_id=session_id,
            filename=safe_name,
            chunk_count=chunk_count,
            message=f"Đã xử lý xong! Tạo {chunk_count} chunks từ '{safe_name}'.",
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=f"Lỗi xử lý: {str(e)}")


@app.post("/api/query", response_model=QueryResponse)
async def query_rag(req: QueryRequest):
    """Hỏi đáp RAG: câu hỏi → tìm chunks → sinh câu trả lời."""
    if req.session_id not in sessions:
        raise HTTPException(
            status_code=400,
            detail="Session không hợp lệ. Vui lòng upload tài liệu trước.",
        )

    question = req.question.strip()
    if not question:
        raise HTTPException(status_code=400, detail="Vui lòng nhập câu hỏi.")

    try:
        session = sessions[req.session_id]
        collection_name = session["collection_name"]

        logger.info(f"Query: {question}")
        relevant_chunks = query_chroma(collection_name, question)

        if not relevant_chunks:
            return QueryResponse(
                answer="Không tìm thấy thông tin liên quan trong tài liệu.",
                sources=[],
            )

        logger.info("Generating answer with Qwen2.5-7B...")
        answer = generate_answer(question, relevant_chunks)

        return QueryResponse(answer=answer, sources=relevant_chunks[:3])

    except Exception as e:
        logger.error(f"Query error: {e}")
        raise HTTPException(status_code=500, detail=f"Lỗi truy vấn: {str(e)}")


@app.get("/api/sessions", response_model=list[SessionInfo])
async def list_sessions():
    """Danh sách các session đang hoạt động."""
    return [
        SessionInfo(
            session_id=sid,
            filename=info["filename"],
            chunk_count=info["chunk_count"],
        )
        for sid, info in sessions.items()
    ]


@app.get("/api/health", response_model=HealthResponse)
async def health_check():
    """Kiểm tra trạng thái server, GPU và models."""
    gpu_available = torch.cuda.is_available()
    gpu_name = torch.cuda.get_device_name(0) if gpu_available else "N/A"
    vram_used = torch.cuda.memory_allocated() / 1e9 if gpu_available else 0
    vram_total = torch.cuda.get_device_properties(0).total_memory / 1e9 if gpu_available else 0
    return HealthResponse(
        status="ok",
        gpu=gpu_available,
        gpu_name=gpu_name,
        vram_used_gb=round(vram_used, 1),
        vram_total_gb=round(vram_total, 1),
        embed_model=EMBED_MODEL_NAME,
        llm_model=LLM_MODEL_NAME,
    )


# ============================================================
# Static files
# ============================================================
app.mount("/static", StaticFiles(directory="static"), name="static")


# ============================================================
# Run: python app.py
# ============================================================
if __name__ == "__main__":
    import uvicorn

    uvicorn.run("app:app", host="0.0.0.0", port=7860, reload=False, log_level="info")